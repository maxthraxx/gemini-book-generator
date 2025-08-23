import hashlib
import logging
import pathlib
import re
import time

import markdown
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt
from lxml import html
from PIL import Image

from book_generator.utils import render_latex_to_image, sanitize_filename

# --- Constants ---
# Multiplier for inline math height based on font size
DEFAULT_INLINE_MATH_HEIGHT_MULTIPLIER = 1.05
# Height in inches for display math images
DEFAULT_DISPLAY_MATH_HEIGHT_INCHES = 0.375


# Helper function to apply formatting to a run
def apply_formatting(run, bold=False, italic=False):
    """Applies formatting to a run."""
    run.bold = bold
    run.italic = italic


# Recursive function to process node content
def process_node_content(
    node,
    paragraph,
    container,
    doc,
    config,
    usable_width_inches,
    equation_image_dir,
    is_bold=False,
    is_italic=False,
):
    """
    Processes an lxml node's content (text and children) recursively,
    adding formatted runs to the paragraph or delegating block elements.
    Inherits formatting state.
    """
    if node.text:
        run = paragraph.add_run(node.text)
        apply_formatting(run, is_bold, is_italic)

    block_tags = {
        "p",
        "ul",
        "ol",
        "table",
        "blockquote",
        "pre",
        "hr",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
        "div",
    }

    for child in node:
        try:
            child_html_snippet = html.tostring(
                child, encoding="unicode", pretty_print=False
            )[:150]
        except Exception:
            child_html_snippet = f"Cannot serialize child <{child.tag}>"
        logging.debug(
            f"Processing child: tag=<{child.tag}>, class='{child.get('class', '')}', "
            f"has_text='{bool(child.text)}', has_tail='{bool(child.tail)}', "
            f"html='{child_html_snippet}...'"
        )

        node_class = child.get("class", "")
        child_tag = child.tag

        new_bold = is_bold or child_tag in ["strong", "b"]
        new_italic = is_italic or child_tag in ["em", "i"]

        is_block = child_tag in block_tags and not (
            child_tag == "div" and "arithmatex" in node_class
        )

        if is_block:
            logging.debug(
                f"-> Encountered block tag <{child_tag}> within inline processing. Calling add_paragraph_from_html_node."
            )
            if child.tail:
                logging.debug(
                    f"   Processing tail of block <{child_tag}> in original paragraph: '{child.tail[:50]}...'"
                )
                run = paragraph.add_run(child.tail)
                apply_formatting(
                    run, is_bold, is_italic
                )  # Apply formatting of the original paragraph context

        elif child_tag == "br":
            run = paragraph.add_run()
            run.add_break()
            if child.tail:
                run = paragraph.add_run(child.tail)
                apply_formatting(run, is_bold, is_italic)

        elif child.tag == "span" and "arithmatex" in node_class:  # Inline Math
            latex_code_with_delimiters = child.text_content().strip()
            logging.debug(
                f"-> Handling Inline Math Span. Raw Content: {latex_code_with_delimiters[:50]}..."
            )

            if latex_code_with_delimiters:
                image_path = render_latex_to_image(
                    latex_code_with_delimiters,
                    equation_image_dir,
                    is_display_style=False,
                )
                if image_path:
                    try:
                        pic_run = paragraph.add_run()
                        font_size_pt = 12
                        if (
                            paragraph.style
                            and paragraph.style.font
                            and paragraph.style.font.size
                        ):
                            font_size_pt = paragraph.style.font.size.pt
                        elif paragraph.runs:
                            for prev_run in reversed(paragraph.runs[:-1]):
                                if prev_run.font and prev_run.font.size:
                                    font_size_pt = prev_run.font.size.pt
                                    logging.debug(
                                        f"Detected font size {font_size_pt}pt from previous run."
                                    )
                                    break
                            else:
                                logging.debug(
                                    "Could not detect font size from previous runs, using default 12pt."
                                )
                        else:
                            logging.debug(
                                "Paragraph has no style/runs with size, using default 12pt."
                            )

                        inline_multiplier = config.get("style_params", {}).get(
                            "inline_math_height_multiplier",
                            DEFAULT_INLINE_MATH_HEIGHT_MULTIPLIER,
                        )
                        calculated_height = Pt(font_size_pt * inline_multiplier)
                        logging.debug(
                            f"Adding picture {image_path} with calculated height {calculated_height}"
                        )
                        pic_run.add_picture(image_path, height=calculated_height)

                        rpr = pic_run._r.get_or_add_rPr()
                        position_element = OxmlElement("w:position")
                        default_offset = -4
                        vertical_offset_half_points = config.get(
                            "style_params", {}
                        ).get("inline_math_vertical_offset_half_points", default_offset)
                        try:
                            vertical_offset_half_points = int(
                                vertical_offset_half_points
                            )
                        except (ValueError, TypeError):
                            logging.warning(
                                f"Invalid value '{vertical_offset_half_points}' for 'inline_math_vertical_offset_half_points' in config. Using default {default_offset}."
                            )
                            vertical_offset_half_points = default_offset
                        position_element.set(
                            qn("w:val"), str(vertical_offset_half_points)
                        )
                        rpr.append(position_element)
                        logging.debug(
                            f"Applied vertical offset ({vertical_offset_half_points} half-points) from config to inline math image run."
                        )

                    except Exception as img_err:
                        logging.error(
                            f"Error adding inline math picture {image_path}: {img_err}"
                        )
                        err_run = paragraph.add_run(
                            f"[Err: Inline Math '{latex_code_with_delimiters[:20]}...']"
                        )
                        apply_formatting(err_run, is_bold, is_italic)
                else:
                    err_run = paragraph.add_run(
                        f"[Render Err: {latex_code_with_delimiters[:20]}...]"
                    )
                    apply_formatting(err_run, is_bold, is_italic)
            else:
                logging.warning("-> Found inline math span but it was empty.")
                err_run = paragraph.add_run("[Err: Empty Math Span]")
                apply_formatting(err_run, is_bold, is_italic)

            if child.tail:
                run = paragraph.add_run(child.tail)
                apply_formatting(run, is_bold, is_italic)

        else:  # General Recursion for other INLINE tags
            logging.debug(f"-> Recursing inline into child <{child.tag}>...")
            process_node_content(
                child,
                paragraph,
                container,
                doc,
                config,
                usable_width_inches,
                equation_image_dir,
                new_bold,
                new_italic,
            )
            if child.tail:
                logging.debug(
                    f"   Processing tail of inline <{child.tag}>: '{child.tail[:50]}...'"
                )
                run = paragraph.add_run(child.tail)
                apply_formatting(run, is_bold, is_italic)


# --- In function: process_mixed_content ---
def process_mixed_content(
    parent_node,
    paragraph,
    container,
    doc,
    config,
    usable_width_inches,
    equation_image_dir,
):
    """
    Starts the recursive processing of mixed content within an HTML node (like p, li, td).
    Adds formatted runs directly to the provided paragraph object or delegates block elements.
    """
    process_node_content(
        parent_node,
        paragraph,
        container,
        doc,
        config,
        usable_width_inches,
        equation_image_dir,
        is_bold=False,
        is_italic=False,
    )


def delete_paragraph(paragraph):
    """Helper function to delete a paragraph."""
    p = paragraph._element
    if p is not None and p.getparent() is not None:
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None


def add_html_content_to_cell(
    html_node, cell, doc, config, usable_width_inches, equation_image_dir
):
    """Adds content from an HTML node (like TD or TH) to a DOCX cell."""
    # Clear existing content (Word adds an empty paragraph by default)
    for p in list(cell.paragraphs):  # Iterate over a copy
        delete_paragraph(p)

    # Add a new paragraph to start processing content
    p = cell.add_paragraph()
    process_mixed_content(
        html_node, p, cell, doc, config, usable_width_inches, equation_image_dir
    )

    # Ensure cell isn't totally empty (Word requires at least one paragraph)
    if not cell.paragraphs or (
        len(cell.paragraphs) == 1 and not cell.paragraphs[0].runs
    ):
        # If processing resulted in no paragraphs or an empty one, ensure one exists.
        if not cell.paragraphs:
            cell.add_paragraph()  # Add a paragraph if none exist
        # If the only paragraph is empty, it's fine, Word needs it.


def add_paragraph_from_html_node(
    node,
    container,
    doc,
    config,
    usable_width_inches,
    equation_image_dir,
    list_level=0,
):
    """
    Processes lxml HTML block nodes and adds them to a python-docx container (doc, cell).
    Handles p, h1-h6, ul, ol, li, blockquote, pre, hr, table, and display math divs/spans within p.
    Delegates inline formatting and nested block handling to process_mixed_content.
    Tracks list nesting depth. Treats all lists (ol and ul) as bulleted lists.
    """
    text = (node.text or "").strip()
    is_handled = False
    node_class = node.get("class", "")

    # --- (Keep add_display_math_image function as is) ---
    def add_display_math_image(image_path, latex_code_for_log):
        # ... (implementation unchanged) ...
        nonlocal container, doc, config, usable_width_inches
        try:
            is_main_body = hasattr(container, "add_table")  # Heuristic check
            p = container.add_paragraph()
            run = p.add_run()

            img_width_inches = None
            img_height_inches = None
            try:
                with Image.open(image_path) as img:
                    width_px, height_px = img.size
                    dpi = 300  # Match render_latex_to_image DPI
                    img_width_inches = width_px / dpi
                    img_height_inches = height_px / dpi
                    logging.debug(
                        f"Image {image_path}: {width_px}x{height_px}px @{dpi}dpi -> {img_width_inches:.2f}x{img_height_inches:.2f} inches"
                    )
            except Exception as pil_err:
                logging.error(
                    f"PIL Error reading image {image_path}: {pil_err}. Cannot determine size for scaling."
                )
                display_height_inches = config.get("style_params", {}).get(
                    "display_math_height_inches", DEFAULT_DISPLAY_MATH_HEIGHT_INCHES
                )
                run.add_picture(image_path, height=Inches(display_height_inches))
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                return

            if (
                is_main_body
                and img_width_inches is not None
                and usable_width_inches is not None
                and img_width_inches > usable_width_inches
            ):
                scale_factor = usable_width_inches / img_width_inches
                final_width_inches = usable_width_inches
                final_height_inches = img_height_inches * scale_factor
                logging.info(
                    f'Scaling display math image {image_path} from {img_width_inches:.2f}" to fit usable width {usable_width_inches:.2f}" (scale: {scale_factor:.2f})'
                )
                run.add_picture(
                    image_path,
                    width=Inches(final_width_inches),
                    height=Inches(final_height_inches),
                )
            elif img_width_inches is not None:
                logging.debug(
                    f'Adding display math image {image_path} ({img_width_inches:.2f}") with original size (fits or not in main body).'
                )
                run.add_picture(
                    image_path,
                    width=Inches(img_width_inches),
                    height=Inches(img_height_inches),
                )
            else:  # Fallback if PIL failed earlier
                display_height_inches = config.get("style_params", {}).get(
                    "display_math_height_inches", DEFAULT_DISPLAY_MATH_HEIGHT_INCHES
                )
                run.add_picture(image_path, height=Inches(display_height_inches))

            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            logging.debug("Display math picture added and centered.")

        except Exception as img_err:
            logging.error(f"Error adding display math picture {image_path}: {img_err}")
            if "p" not in locals():
                p = container.add_paragraph()  # Ensure p exists
            p.add_run(
                f"[Error adding display math image: {latex_code_for_log[:30]}...]"
            )
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Handle Display Math (Arithmatex Div) ---
    if node.tag == "div" and "arithmatex" in node_class:
        # ... (implementation unchanged) ...
        logging.debug(
            f"Processing display math DIV node: <{node.tag} class='{node_class}'>"
        )
        latex_code_with_delimiters = node.text_content().strip()
        if latex_code_with_delimiters:
            image_path = render_latex_to_image(
                latex_code_with_delimiters, equation_image_dir, is_display_style=True
            )
            if image_path:
                logging.debug(f"Successfully rendered display math to {image_path}")
                add_display_math_image(image_path, latex_code_with_delimiters)
            else:
                logging.warning(
                    f"Rendering failed for display math: {latex_code_with_delimiters[:50]}..."
                )
                if hasattr(container, "add_paragraph"):
                    p = container.add_paragraph(
                        f"[Render Err: Display Math {latex_code_with_delimiters[:30]}...]"
                    )
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            logging.warning("Found display math div but it was empty.")
        is_handled = True

    # --- Handle Paragraphs (<p>) ---
    elif node.tag == "p":
        # ... (implementation unchanged, including display math span detection) ...
        logging.debug(f"Processing <p> node.")
        children = list(node)
        # Check if <p> *only* contains a display math span (common Arithmatex output)
        if (
            not text  # No text directly in <p>
            and len(children) == 1
            and children[0].tag == "span"
            and "arithmatex" in children[0].get("class", "")
            and not (children[0].tail or "").strip()  # No tail text after span
        ):
            logging.debug(
                "Detected <p> containing only an arithmatex span. Treating as display math."
            )
            span_node = children[0]
            latex_code_with_delimiters = span_node.text_content().strip()
            if latex_code_with_delimiters:
                image_path = render_latex_to_image(
                    latex_code_with_delimiters,
                    equation_image_dir,
                    is_display_style=True,
                )
                if image_path:
                    add_display_math_image(image_path, latex_code_with_delimiters)
                else:
                    logging.warning(
                        f"Rendering failed for display math span in p: {latex_code_with_delimiters[:50]}..."
                    )
                    if hasattr(container, "add_paragraph"):
                        p = container.add_paragraph(
                            f"[Render Err: Display Math {latex_code_with_delimiters[:30]}...]"
                        )
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                logging.warning("Found display math span within p but it was empty.")
            is_handled = True

        # --- Default paragraph handling ---
        if not is_handled:
            logging.debug("  Processing <p> using default process_mixed_content.")
            current_paragraph = container.add_paragraph()
            process_mixed_content(
                node,
                current_paragraph,
                container,
                doc,
                config,
                usable_width_inches,
                equation_image_dir,
            )
            # Remove paragraph if it ended up empty after processing
            if not current_paragraph.text and not current_paragraph.runs:
                logging.debug("Removing empty paragraph added for <p>.")
                delete_paragraph(current_paragraph)

            is_handled = True

    # --- Handle Tables ---
    elif node.tag == "table":
        # ... (implementation unchanged) ...
        logging.debug("Processing table...")
        html_rows = node.xpath(".//tr")  # Get all rows in the table
        if not html_rows:
            logging.warning("Table tag found but contains no rows (tr). Skipping.")
        else:
            # Determine number of columns from the first row
            first_row_cells = html_rows[0].xpath("./th|./td")
            num_cols = len(first_row_cells)
            if num_cols == 0:
                logging.warning(
                    "Table's first row contains no cells (th/td). Skipping table."
                )
            else:
                # Add table to the container (doc or cell)
                docx_table = container.add_table(rows=0, cols=num_cols)
                docx_table.style = "Table Grid"  # Apply a basic style

                # Process each row
                for html_row in html_rows:
                    docx_row = docx_table.add_row()
                    html_cells = html_row.xpath("./th|./td")
                    # Process each cell in the row
                    for i, cell_node in enumerate(html_cells):
                        if (
                            i < num_cols
                        ):  # Avoid index errors if rows have varying cell counts
                            docx_cell = docx_row.cells[i]
                            # Use the dedicated function to populate the cell
                            add_html_content_to_cell(
                                cell_node,
                                docx_cell,
                                doc,
                                config,
                                usable_width_inches,
                                equation_image_dir,
                            )
                        else:
                            logging.warning(
                                f"Row has more cells ({len(html_cells)}) than table columns ({num_cols}). Ignoring extra cells."
                            )
        is_handled = True

    # --- Handle Headings (h1-h6) ---
    elif node.tag in ["h1", "h2", "h3", "h4", "h5", "h6"]:
        # ... (implementation unchanged) ...
        level = int(node.tag[1])
        style_name = f"Heading {level}"
        # Use default paragraph style if heading style doesn't exist
        style = (
            doc.styles[style_name] if style_name in doc.styles else doc.styles["Normal"]
        )
        p = container.add_paragraph(style=style)
        process_mixed_content(
            node, p, container, doc, config, usable_width_inches, equation_image_dir
        )
        is_handled = True

    # --- Handle Lists (ul, ol, li) ---
    elif node.tag in ["ul", "ol"]:
        # Keep track of the level for children of THIS list
        child_list_level = list_level + 1
        logging.debug(
            f"Processing <{node.tag}> at list_level {list_level}. Children will be level {child_list_level}."
        )
        for child_node in node:
            # Only process li, ul, ol directly. Ignore whitespace text nodes etc.
            if child_node.tag == "li":
                # Pass the PARENT list's level down to the li
                add_paragraph_from_html_node(
                    child_node,
                    container,
                    doc,
                    config,
                    usable_width_inches,
                    equation_image_dir,
                    list_level=list_level,  # Pass current level to li
                )
            elif child_node.tag in ["ul", "ol"]:
                # Handle potentially invalid nested lists directly under lists
                logging.warning(
                    f"Found nested <{child_node.tag}> directly inside <{node.tag}>. Processing recursively."
                )
                add_paragraph_from_html_node(
                    child_node,
                    container,
                    doc,
                    config,
                    usable_width_inches,
                    equation_image_dir,
                    list_level=child_list_level,  # Increment level here
                )
            # Ignore other tags or text directly within ul/ol for now, or log warnings
            elif (child_node.text or "").strip():
                logging.warning(
                    f"Ignoring text '{child_node.text.strip()[:50]}...' found directly inside <{node.tag}>."
                )

        is_handled = True

    elif node.tag == "li":
        # Determine the style based on the level of the list this li belongs to.
        # Level 0 -> "List Bullet", Level 1 -> "List Bullet 2", etc.
        current_level_index = list_level + 1  # 1-based index for style name
        style_suffix = f" {current_level_index}" if current_level_index > 1 else ""
        # Assuming bullet lists for now, adapt if numbered lists needed different base
        style_name = f"List Bullet{style_suffix}"

        if style_name not in doc.styles:
            logging.warning(
                f"Style '{style_name}' not found. Falling back to 'List Bullet' or 'Normal'."
            )
            # Fallback logic: try base list style, then Normal
            style_name = "List Bullet" if "List Bullet" in doc.styles else "Normal"

        style = doc.styles[style_name]
        logging.debug(
            f"Processing <li> at list_level {list_level} using style '{style.name}'."
        )

        # Create the paragraph for this list item's text/inline content
        # Check if the first child is a <p> tag (common markdown output)
        first_child_is_p = False
        direct_children = [child for child in node if isinstance(child.tag, str)]
        if (
            direct_children
            and direct_children[0].tag == "p"
            and not (node.text or "").strip()
        ):
            # If li starts directly with a <p>, use its content but apply list style
            p_node = direct_children[0]
            p = container.add_paragraph(style=style)
            process_mixed_content(
                p_node,
                p,
                container,
                doc,
                config,
                usable_width_inches,
                equation_image_dir,
            )
            # Remove the processed <p> node from children list to avoid double processing
            children_to_process = direct_children[1:]
            logging.debug("Processed <li> content starting with <p>.")
        else:
            # Process text/inline elements directly under <li>
            p = container.add_paragraph(style=style)
            process_mixed_content(
                node, p, container, doc, config, usable_width_inches, equation_image_dir
            )
            # We processed the whole node inline content, but need to handle block children (nested lists) separately
            children_to_process = (
                direct_children  # Re-evaluate children for nested lists
            )
            logging.debug("Processed <li> content directly.")

        # Now, specifically look for and handle nested lists *within* this <li>
        nested_list_found = False
        for (
            child_node
        ) in children_to_process:  # Use the adjusted list if first child was <p>
            if child_node.tag in ["ul", "ol"]:
                nested_list_found = True
                # Recursively call for the nested list, INCREMENTING the level
                nested_list_level = list_level + 1
                logging.debug(
                    f"  Found nested <{child_node.tag}> inside <li>. Processing at level {nested_list_level}."
                )
                add_paragraph_from_html_node(
                    child_node,
                    container,
                    doc,
                    config,
                    usable_width_inches,
                    equation_image_dir,
                    list_level=nested_list_level,  # Pass incremented level
                )
            # Note: We assume process_mixed_content handled inline tags and text already.
            # If there were other block elements inside <li> (besides <p> handled above),
            # they might need specific handling here too, but nested lists are primary.

        # Remove the list item paragraph ONLY if it's empty AND no nested list followed.
        # Check runs as well as text, because an image (like math) adds a run but no text.
        if not p.text.strip() and not p.runs and not nested_list_found:
            logging.debug(
                f"Removing empty paragraph potentially created for <li> at level {list_level}."
            )
            delete_paragraph(p)

        is_handled = True

    # --- Handle Blockquotes ---
    elif node.tag == "blockquote":
        logging.debug("Processing <blockquote> node.")
        style = doc.styles["Quote"] if "Quote" in doc.styles else doc.styles["Normal"]
        for child_node in node:
            if child_node.tag == "p":
                p = container.add_paragraph(style=style)
                process_mixed_content(
                    child_node,
                    p,
                    container,
                    doc,
                    config,
                    usable_width_inches,
                    equation_image_dir,
                )
                if not p.text and not p.runs:
                    logging.debug(
                        "Removing empty paragraph added for <p> inside <blockquote>."
                    )
                    delete_paragraph(p)
            elif child_node.tag is not None:
                logging.debug(
                    f"Handling non-<p> tag '{child_node.tag}' inside blockquote."
                )
                add_paragraph_from_html_node(
                    child_node,
                    container,
                    doc,
                    config,
                    usable_width_inches,
                    equation_image_dir,
                )
            elif (child_node.text or "").strip():
                logging.debug("Handling text node directly inside <blockquote>.")
                p = container.add_paragraph(child_node.text.strip(), style=style)
        is_handled = True

    # --- Handle Highlighted Code Blocks (div.highlight > pre > code) ---
    elif node.tag == "div" and "highlight" in node.get("class", "").split():
        # ... (implementation unchanged) ...
        logging.debug("Processing <div class='highlight'> node.")
        pre_node = node.find("pre")
        if pre_node is not None:
            code_node = pre_node.find("code")
            # Extract text from <code> if present, otherwise from <pre>
            full_text = (
                "".join(pre_node.itertext())
                if code_node is None
                else "".join(code_node.itertext())
            )
            if full_text:
                code_style_name = "CodeBlock"
                style_to_use = (
                    doc.styles[code_style_name]
                    if code_style_name in doc.styles
                    else doc.styles["Normal"]
                )
                logging.debug(
                    f"Applying style '{style_to_use.name}' to <div class='highlight'> content."
                )

                p = container.add_paragraph(full_text.strip("\n"), style=style_to_use)

                # Explicitly set font on runs as a fallback
                for run in p.runs:
                    if not run.font.name or run.font.name != "Courier New":
                        run.font.name = "Courier New"
                    if not run.font.size or run.font.size != Pt(10):
                        run.font.size = Pt(10)
                logging.debug(
                    "Ensured Courier New/10pt font on runs within the highlight code block."
                )
            else:
                logging.debug("Highlight div found, but contained no text.")
        else:
            logging.warning(
                "Found <div class='highlight'> but no <pre> tag inside. Skipping."
            )
        is_handled = True

    # --- Handle Preformatted Text (<pre>) ---
    elif node.tag == "pre":
        # ... (implementation unchanged) ...
        parent = node.getparent()
        if (
            parent is not None
            and parent.tag == "div"
            and "highlight" in parent.get("class", "").split()
        ):
            logging.debug(
                "Skipping <pre> inside already handled <div class='highlight'>."
            )
            is_handled = True
        else:
            logging.debug("Processing plain <pre> node (not inside highlight div).")
            code_node = node.find("code")
            full_text = (
                "".join(node.itertext())
                if code_node is None
                else "".join(code_node.itertext())
            )
            if full_text:
                code_style_name = "CodeBlock"
                style_to_use = (
                    doc.styles[code_style_name]
                    if code_style_name in doc.styles
                    else doc.styles["Normal"]
                )
                logging.debug(f"Applying style '{style_to_use.name}' to <pre> content.")

                p = container.add_paragraph(full_text.strip("\n"), style=style_to_use)

                for run in p.runs:
                    if not run.font.name or run.font.name != "Courier New":
                        run.font.name = "Courier New"
                    if not run.font.size or run.font.size != Pt(10):
                        run.font.size = Pt(10)
                logging.debug(
                    "Ensured Courier New/10pt font on runs within the code block."
                )
            is_handled = True

    # --- Handle Horizontal Rule (<hr>) ---
    elif node.tag == "hr":
        # ... (implementation unchanged) ...
        logging.debug("Processing <hr> node.")
        p = container.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        pPr.insert_element_before(
            pBdr,
            "w:shd",
            "w:tabs",
            "w:suppressAutoHyphens",
            "w:kinsoku",
            "w:wordWrap",
            "w:overflowPunct",
            "w:topLinePunct",
            "w:autoSpaceDE",
            "w:autoSpaceDN",
            "w:bidi",
            "w:adjustRightInd",
            "w:snapToGrid",
            "w:spacing",
            "w:ind",
            "w:contextualSpacing",
            "w:mirrorIndents",
            "w:suppressOverlap",
            "w:jc",
            "w:textDirection",
            "w:textAlignment",
            "w:textboxTightWrap",
            "w:outlineLvl",
            "w:divId",
            "w:cnfStyle",
            "w:rPr",
            "w:sectPr",
            "w:pPrChange",
        )
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")  # 3/4 pt
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom)
        is_handled = True

    # --- Fallback for Unhandled Block Tags ---
    known_inline_or_handled = {
        "strong",
        "b",
        "em",
        "i",
        "span",
        "br",
        "a",
        "code",  # Common inline
        "p",
        "ul",
        "ol",
        "li",
        "table",
        "tr",
        "td",
        "th",  # Handled block
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",  # Handled block
        "blockquote",
        "pre",
        "hr",
        "div",  # Handled block (incl. arithmatex)
    }
    if (
        not is_handled
        and node.tag is not None  # Ensure it's a tag, not comment/text
        and node.tag not in known_inline_or_handled
    ):
        logging.warning(
            f"Unhandled block tag <{node.tag}> encountered. Attempting to process its content as plain text."
        )
        plain_text = "".join(node.itertext()).strip()
        if plain_text:
            container.add_paragraph(plain_text)
            logging.debug(
                f"Added text content of unhandled <{node.tag}>: '{plain_text[:100]}...'"
            )
        else:
            logging.debug(f"Unhandled block tag <{node.tag}> had no text content.")
        is_handled = True


def markdown_to_docx(
    markdown_text,
    container_obj,
    doc,
    config,
    usable_width_inches,
    equation_image_dir,
    context_label=None,
):
    """Converts Markdown to DOCX elements, using Arithmatex for LaTeX,
    and cleans up excessive line breaks. Handles nested lists."""
    # ... (keep markdown cleaning and HTML conversion as is) ...
    if not markdown_text:
        logging.debug("Markdown text is empty, skipping conversion.")
        return

    logging.debug("Starting Markdown to DOCX conversion (using Arithmatex)...")

    # --- Pre-processing: Clean up excessive newlines ---
    cleaned_markdown = re.sub(r"\n{3,}", "\n\n", markdown_text.strip())
    if cleaned_markdown != markdown_text.strip():
        logging.debug("Cleaned excessive newlines from Markdown content.")
    # --- End Pre-processing ---

    try:
        # Configure Markdown extensions
        extensions = [
            "extra",  # Includes tables, footnotes, abbreviations, etc.
            "sane_lists",
            "fenced_code",
            "pymdownx.arithmatex",  # For LaTeX math $$...$$ and $...$
            "pymdownx.superfences",  # Improved fenced code blocks
            "pymdownx.details",
            "pymdownx.mark",
        ]

        extension_configs = {"pymdownx.arithmatex": {"generic": True}}
        md_converter = markdown.Markdown(
            extensions=extensions, extension_configs=extension_configs
        )
        html_content = md_converter.convert(cleaned_markdown)
        logging.debug(f"Generated HTML (first 500 chars): {html_content[:500]}...")

        # --- Save intermediate HTML for debugging ---
        if config.get("debug_options", {}).get("save_intermediate_html", False):
            # ... (HTML saving logic unchanged) ...
            debug_html_dir = (
                pathlib.Path(config.get("cache_dir", "api_cache")) / "debug_html"
            )
            debug_html_dir.mkdir(parents=True, exist_ok=True)
            timestamp = time.strftime("%Y%m%d_%H%M%S")
            content_hash = hashlib.sha1(cleaned_markdown.encode("utf-8")).hexdigest()[
                :8
            ]
            if context_label:
                context_hint = sanitize_filename(context_label, max_length=80)
            else:
                context_hint = "unknown_context"

            html_filename = (
                debug_html_dir
                / f"md_to_html_{context_hint}_{timestamp}_{content_hash}.html"
            )
            try:
                with open(html_filename, "w", encoding="utf-8") as f_html:
                    f_html.write(
                        "<!DOCTYPE html>\n<html>\n<head><meta charset='UTF-8'>"
                    )
                    f_html.write(
                        "<style>.arithmatex { border: 1px dotted blue; padding: 2px; }</style>"
                    )
                    f_html.write("</head>\n<body>\n")
                    f_html.write(html_content)
                    f_html.write("\n</body>\n</html>")
                logging.info(f"Saved intermediate HTML to: {html_filename}")
            except Exception as e_save:
                logging.error(
                    f"Could not save intermediate HTML to {html_filename}: {e_save}"
                )
        # --- End intermediate HTML saving ---

        # --- HTML Cleaning Step ---
        parser = html.HTMLParser(encoding="utf-8")
        # Wrap in a div to ensure a single root for parsing potentially fragmented HTML
        html_wrapper_str = f"<div>{html_content}</div>"
        try:
            # Use utf-8 encoding for parsing
            html_tree_root = html.fromstring(
                html_wrapper_str.encode("utf-8"), parser=parser
            )
        except UnicodeDecodeError:
            # Fallback if utf-8 fails (less common but possible)
            logging.warning(
                "UTF-8 decoding failed for HTML string, trying 'latin-1' fallback."
            )
            html_tree_root = html.fromstring(
                html_wrapper_str.encode("latin-1"), parser=parser
            )

        # Find <p> tags that are either completely empty or contain only <br> tags and whitespace
        paragraphs_to_remove = []
        # Iterate through all <p> tags in the parsed HTML tree
        for p_tag in html_tree_root.xpath(".//p"):
            # Get the text content of the <p> tag, stripping leading/trailing whitespace
            text_content = p_tag.text_content().strip()
            # Get all direct children elements of the <p> tag
            children = p_tag.getchildren()

            # Check condition 1: Is the paragraph completely empty (no text, no children)?
            is_completely_empty = not text_content and not children

            # Check condition 2: Does the paragraph contain ONLY <br> tags (and whitespace)?
            only_br_children = False
            if not text_content and children:  # Only check children if there's no text
                all_children_are_br = True
                for child in children:
                    # If any child is not a <br> tag, this condition is false
                    if child.tag != "br":
                        all_children_are_br = False
                        break
                only_br_children = all_children_are_br  # True if all children were <br>

            # If either condition is met, schedule the paragraph for removal
            if is_completely_empty or only_br_children:
                logging.debug(
                    f"Found <p> tag {'empty' if is_completely_empty else 'containing only <br> tags'}. Scheduling for removal. "
                    f"HTML snippet: {html.tostring(p_tag, encoding='unicode', pretty_print=False)[:100]}"
                )
                paragraphs_to_remove.append(p_tag)

        # Remove the identified paragraphs
        if paragraphs_to_remove:
            logging.info(
                f"Removing {len(paragraphs_to_remove)} empty or <br>-only <p> tags."
            )
            for p_tag in paragraphs_to_remove:
                parent = p_tag.getparent()
                if parent is not None:
                    # Preserve tail text if it exists, attaching it to the previous sibling or parent text
                    if p_tag.tail and p_tag.tail.strip():
                        previous_sibling = p_tag.getprevious()
                        if previous_sibling is not None:
                            # Append tail to the previous sibling's tail
                            previous_sibling.tail = (
                                previous_sibling.tail or ""
                            ) + p_tag.tail
                        else:
                            # Append tail to the parent's text if no previous sibling
                            parent.text = (parent.text or "") + p_tag.tail
                    # Remove the paragraph tag itself
                    parent.remove(p_tag)
        # --- End HTML Cleaning Step ---

        # Process each *remaining* element under the root div
        for element in html_tree_root:
            add_paragraph_from_html_node(
                element,
                container_obj,
                doc,
                config,
                usable_width_inches,
                equation_image_dir,
            )

    except ImportError as ie:
        # ... (error handling unchanged) ...
        logging.error(
            f"Markdown extension import error: {ie}. Ensure required libraries (e.g., pymdown-extensions) are installed.",
            exc_info=True,
        )
        container_obj.add_paragraph(f"[Error: Missing Markdown extension - {ie}]")
        container_obj.add_paragraph(cleaned_markdown)
    except Exception as e:
        # ... (error handling unchanged) ...
        logging.error(f"Error converting Markdown/HTML to DOCX: {e}", exc_info=True)
        container_obj.add_paragraph(f"[Error processing content: {e}]")
        container_obj.add_paragraph("--- Raw Markdown Fallback ---")
        container_obj.add_paragraph(cleaned_markdown)
        container_obj.add_paragraph("--- End Raw Markdown ---")

    logging.debug("Finished Markdown to DOCX conversion.")


def set_page_numbering(section, format_code, start_number=None):
    """Adds page numbering to the footer of a given section."""
    # Ensure footer exists (it might not by default)
    if section.footer is None:
        section.footer  # Accessing it creates it
        logging.debug("Created footer for section.")

    footer = section.footer

    # Unlink footer if setting a start number and it was linked
    if footer.is_linked_to_previous and start_number is not None:
        footer.is_linked_to_previous = False
        logging.debug(f"Unlinking footer for section starting page {start_number}")

    # Clear existing footer content (if any) before adding page number field
    if footer.paragraphs:
        pf = footer.paragraphs[0]
        pf.clear()  # Clear runs from the first paragraph
        # Remove other paragraphs if they exist
        for p in footer.paragraphs[1:]:
            delete_paragraph(p)
    else:
        # Add a paragraph if none exist
        pf = footer.add_paragraph()

    pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run_begin = pf.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_char_begin)

    run_instr = pf.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    run_instr._r.append(instr_text)

    run_sep = pf.add_run()
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run_sep._r.append(fld_char_separate)

    # Optional: Add a run here for the actual number display if needed,
    # but Word usually handles this automatically with the fields.

    run_end = pf.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_char_end)
    # --- End PAGE field ---

    # --- Set page number type in section properties ---
    sectPr = section._sectPr
    # Remove existing pgNumType if it exists to avoid conflicts
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is not None:
        sectPr.remove(pgNumType)

    # Create and append new pgNumType element
    pgNumType = OxmlElement("w:pgNumType")
    sectPr.append(pgNumType)  # Append to the end or specific location if needed

    # Set format
    pgNumType.set(qn("w:fmt"), format_code)

    # Set start number if provided
    if start_number is not None:
        pgNumType.set(qn("w:start"), str(start_number))
        logging.debug(
            f"Set page numbering: format='{format_code}', start={start_number}"
        )
    else:
        # Ensure start attribute is removed if not specified (to allow continuation)
        if qn("w:start") in pgNumType.attrib:
            del pgNumType.attrib[qn("w:start")]
        logging.debug(
            f"Set page numbering: format='{format_code}', continuing from previous section."
        )


def assemble_docx(
    config,
    front_matter,
    body_matter,
    back_matter,
    book_title,
    equation_image_dir,
    output_dir,
):
    """Assembles the main book DOCX file with complex page numbering and MathML/OXML."""
    logging.info("Assembling main DOCX file...")
    filename_stem = sanitize_filename(book_title)
    # Construct the full output path using the provided directory
    output_filename = output_dir / f"{filename_stem}.docx"
    logging.info(f"Main book filename set to: '{output_filename}'")

    style_config = config.get("style_params", {})
    font_name = style_config.get("font_name", "Times New Roman")
    font_size = style_config.get("font_size", 12)
    page_preset = style_config.get("page_size_preset", "6x9")

    margin_config_mm = style_config.get("margins_mm", {})
    default_top_mm = 19  # Approx 0.75 inch
    default_bottom_mm = 19
    default_inside_mm = 19  # For gutter
    default_outside_mm = 13  # Approx 0.5 inch
    default_gutter_mm = 0  # Set gutter explicitly

    top_margin_mm = margin_config_mm.get("top", default_top_mm)
    bottom_margin_mm = margin_config_mm.get("bottom", default_bottom_mm)
    # 'left' in config maps to 'outside', 'right' maps to 'inside' for mirrored margins
    outside_margin_mm = margin_config_mm.get("left", default_outside_mm)
    inside_margin_mm = margin_config_mm.get("right", default_inside_mm)
    gutter_margin_mm = margin_config_mm.get("gutter", default_gutter_mm)

    doc = Document()

    # --- Basic Style Setup ---
    try:
        style = doc.styles["Normal"]
        style.font.name = font_name
        style.font.size = Pt(font_size)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # Ensure heading styles use the base font and add spacing
        for i in range(1, 7):
            heading_style_name = f"Heading {i}"
            if heading_style_name in doc.styles:
                h_style = doc.styles[heading_style_name]
                h_style.font.name = font_name
                # Add some default spacing (can be overridden in config later)
                if i == 1:
                    h_style.paragraph_format.space_before = Pt(18)
                    h_style.paragraph_format.space_after = Pt(6)
                elif i == 2:
                    h_style.paragraph_format.space_before = Pt(12)
                    h_style.paragraph_format.space_after = Pt(4)
                else:
                    h_style.paragraph_format.space_before = Pt(6)
                    h_style.paragraph_format.space_after = Pt(2)

        # Ensure Title/Subtitle styles use the base font (or define them)
        if "Title" not in doc.styles:
            title_style = doc.styles.add_style("Title", WD_STYLE_TYPE.PARAGRAPH)
            title_style.base_style = doc.styles["Normal"]
            title_style.font.name = font_name
            title_style.font.size = Pt(28)  # Example size
            title_style.font.bold = True
            title_style.paragraph_format.space_after = Pt(6)
            title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            doc.styles["Title"].font.name = font_name
            doc.styles["Title"].paragraph_format.alignment = (
                WD_PARAGRAPH_ALIGNMENT.CENTER
            )

        if "Subtitle" not in doc.styles:
            subtitle_style = doc.styles.add_style("Subtitle", WD_STYLE_TYPE.PARAGRAPH)
            subtitle_style.base_style = doc.styles["Normal"]
            subtitle_style.font.name = font_name
            subtitle_style.font.size = Pt(16)  # Example size
            subtitle_style.font.italic = True
            subtitle_style.paragraph_format.space_after = Pt(18)
            subtitle_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            doc.styles["Subtitle"].font.name = font_name
            doc.styles["Subtitle"].paragraph_format.alignment = (
                WD_PARAGRAPH_ALIGNMENT.CENTER
            )

        # Ensure List Bullet exists (as before)
        if "List Bullet" not in doc.styles:
            lb_style = doc.styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)
            lb_style.base_style = doc.styles["Normal"]
            # TODO: Define actual bullet point and indentation via Oxml or basic properties
            lb_style.paragraph_format.left_indent = Inches(0.25)  # Example
            lb_style.paragraph_format.first_line_indent = Inches(
                -0.25
            )  # Example hanging indent
        else:  # Ensure base style has some indent
            lb_style = doc.styles["List Bullet"]
            if lb_style.paragraph_format.left_indent is None:
                lb_style.paragraph_format.left_indent = Inches(0.25)
            if lb_style.paragraph_format.first_line_indent is None:
                lb_style.paragraph_format.first_line_indent = Inches(-0.25)

        # Define nested styles (add more as needed)
        for i in range(2, 5):  # Define List Bullet 2, 3, 4
            style_name = f"List Bullet {i}"
            base_style_name = (
                f"List Bullet {i-1}" if i > 2 else "List Bullet"
            )  # Base on previous level
            if style_name not in doc.styles:
                lb_nested_style = doc.styles.add_style(
                    style_name, WD_STYLE_TYPE.PARAGRAPH
                )
                # Base on previous level if possible, otherwise Normal
                base_style = (
                    doc.styles[base_style_name]
                    if base_style_name in doc.styles
                    else doc.styles["Normal"]
                )
                lb_nested_style.base_style = base_style
                # Increase indentation relative to base style or set absolute
                # Example: Add 0.25 inches per level
                indent_inches = (i - 1) * 0.35  # Adjust multiplier as needed
                lb_nested_style.paragraph_format.left_indent = Inches(indent_inches)
                # Keep hanging indent consistent or adjust if needed
                lb_nested_style.paragraph_format.first_line_indent = Inches(-0.25)
                logging.info(
                    f"Defined style '{style_name}' with left indent {indent_inches} inches."
                )
            else:
                # Optionally ensure indentation is correct on existing styles
                lb_nested_style = doc.styles[style_name]
                indent_inches = (i - 1) * 0.35
                if lb_nested_style.paragraph_format.left_indent != Inches(
                    indent_inches
                ):
                    logging.debug(f"Adjusting indent for existing style '{style_name}'")
                    lb_nested_style.paragraph_format.left_indent = Inches(indent_inches)
                    lb_nested_style.paragraph_format.first_line_indent = Inches(-0.25)

        # --- Define Code Block Style ---
        code_style_name = "CodeBlock"
        if code_style_name not in doc.styles:
            code_style = doc.styles.add_style(code_style_name, WD_STYLE_TYPE.PARAGRAPH)
            # Base on 'No Spacing' if it exists for minimal vertical space, else 'Normal'
            base_style_name = "No Spacing" if "No Spacing" in doc.styles else "Normal"
            code_style.base_style = doc.styles[base_style_name]
            code_style.font.name = "Courier New"  # Monospace font
            code_style.font.size = Pt(10)  # Slightly smaller size often looks good
            # Optional: Add indentation or borders
            # code_style.paragraph_format.left_indent = Inches(0.25)
            # Optional: Adjust spacing if needed (base style might handle it)
            # code_style.paragraph_format.space_before = Pt(6)
            # code_style.paragraph_format.space_after = Pt(6)
            logging.info(
                f"Defined '{code_style_name}' style based on '{base_style_name}'."
            )
        else:
            # Ensure existing style uses monospace font
            existing_code_style = doc.styles[code_style_name]
            existing_code_style.font.name = "Courier New"
            existing_code_style.font.size = Pt(10)
            logging.info(f"Ensured '{code_style_name}' style uses Courier New, 10pt.")
        # --- End Code Block Style ---

        logging.info("Styles configured.")
    except Exception as e:
        logging.error(f"Error setting up styles: {e}")

    # --- Page Setup (Initial Section - Section 0) ---
    section0 = doc.sections[0]
    page_width_mm = None

    if page_preset == "6x9":
        section0.page_width = Inches(6)
        section0.page_height = Inches(9)
        logging.info("Set page size to 6x9 inches.")
    elif page_preset == "A4":
        section0.page_width = Mm(210)
        section0.page_height = Mm(297)
        logging.info("Set page size to A4.")
    else:
        logging.warning(
            f"Unsupported page_size_preset '{page_preset}'. Using default Word size."
        )
        # Use default size implicitly

    if section0.page_width is not None:
        page_width_mm = section0.page_width / Mm(1)
        logging.info(f"Actual page width from section object: {page_width_mm:.2f} mm")
    else:
        logging.warning("Could not determine page width from section object.")

    # --- Apply Margins and Gutter (Mirrored) ---
    try:
        section0.top_margin = Mm(top_margin_mm)
        section0.bottom_margin = Mm(bottom_margin_mm)
        section0.left_margin = Mm(outside_margin_mm)  # Outside margin
        section0.right_margin = Mm(inside_margin_mm)  # Inside margin
        section0.gutter = Mm(gutter_margin_mm)
        # Enable mirrored margins for gutter to work correctly
        sectPr = section0._sectPr
        pgMar = sectPr.find(qn("w:pgMar"))
        if pgMar is not None:
            pgMar.set(qn("w:mirrorMargins"), "true")  # Use 'true' or '1'
        logging.info(
            f"Set mirrored margins (mm): Top={top_margin_mm}, Bottom={bottom_margin_mm}, "
            f"Outside={outside_margin_mm}, Inside={inside_margin_mm}, Gutter={gutter_margin_mm}"
        )
    except ValueError as ve:
        logging.error(f"Invalid margin value provided: {ve}. Using Word defaults.")
    except Exception as e:
        logging.error(f"Error setting margins: {e}. Using Word defaults.")

    # --- Calculate Usable Width ---
    usable_width_mm = None
    usable_width_inches = None
    if page_width_mm is not None:
        # Usable width = Page Width - Outside Margin - Inside Margin - Gutter
        usable_width_mm = (
            page_width_mm - outside_margin_mm - inside_margin_mm - gutter_margin_mm
        )
        usable_width_inches = usable_width_mm / 25.4
        logging.info(
            f"Calculated usable page width: {usable_width_mm:.2f} mm ({usable_width_inches:.2f} inches)"
        )
    else:
        logging.error("Cannot calculate usable width because page width is unknown.")
    # --- End Calculate Usable Width ---

    # --- Section 0: Title Page ---
    logging.info("Adding Title Page (Section 0)...")
    if "title_page" in front_matter:
        tp_info = front_matter["title_page"]
        # Add space before title (adjust as needed)
        doc.add_paragraph().paragraph_format.space_before = Pt(72)

        title_p = doc.add_paragraph(tp_info["title"], style="Title")
        # Alignment is set in style definition now

        if tp_info.get("subtitle"):
            subtitle_p = doc.add_paragraph(tp_info["subtitle"], style="Subtitle")
            # Alignment and spacing set in style definition
        else:
            # Add extra space after title if no subtitle
            title_p.paragraph_format.space_after = Pt(
                24
            )  # Override style default if needed

        # Add Author with space before it
        author_p = doc.add_paragraph()  # Empty paragraph for spacing
        author_p.paragraph_format.space_before = Pt(36)  # Space before author
        author_p.add_run(f"By {tp_info['author']}")
        author_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Optional: Create and apply an 'Author' style

    # --- Section Break for Copyright Page (Starts Section 1) ---
    doc.add_section(WD_SECTION.ODD_PAGE)
    section1 = doc.sections[1]
    # Copy page setup from section 0 to section 1
    section1.page_height = section0.page_height
    section1.page_width = section0.page_width
    section1.left_margin = section0.left_margin  # Outside
    section1.right_margin = section0.right_margin  # Inside
    section1.top_margin = section0.top_margin
    section1.bottom_margin = section0.bottom_margin
    section1.gutter = section0.gutter
    # Ensure mirrored margins are also copied
    sectPr1 = section1._sectPr
    pgMar1 = sectPr1.find(qn("w:pgMar"))
    if pgMar1 is not None:
        pgMar1.set(qn("w:mirrorMargins"), "true")
    # No page numbering for title (section 0) or copyright (section 1)

    # --- Section 1: Copyright Page ---
    logging.info("Adding Copyright Page (Section 1)...")
    if "copyright_page" in front_matter:
        cp_text = front_matter["copyright_page"]
        # Split into paragraphs based on double line breaks in the original string
        cp_paragraphs = re.split(r"\n\s*\n", cp_text)

        # Add space before the first paragraph
        first_cp_p = doc.add_paragraph()
        first_cp_p.paragraph_format.space_before = Pt(60)

        # FIX: Replace internal newlines with spaces before adding
        first_para_text = re.sub(r"\s*\n\s*", " ", cp_paragraphs[0].strip())
        first_cp_p.add_run(first_para_text)

        first_cp_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # Apply smaller font size
        for run in first_cp_p.runs:
            run.font.size = Pt(font_size - 2)  # Assuming font_size is defined

        # Add subsequent paragraphs
        for cp_para in cp_paragraphs[1:]:
            # FIX: Replace internal newlines with spaces before adding
            para_text = re.sub(r"\s*\n\s*", " ", cp_para.strip())
            p = doc.add_paragraph(para_text)

            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for run in p.runs:
                run.font.size = Pt(font_size - 2)  # Assuming font_size is defined

    # --- Section Break for Rest of Front Matter (Starts Section 2) ---
    doc.add_section(WD_SECTION.ODD_PAGE)
    section2 = doc.sections[2]
    # Copy page setup from section 0 to section 2
    section2.page_height = section0.page_height
    section2.page_width = section0.page_width
    section2.left_margin = section0.left_margin  # Outside
    section2.right_margin = section0.right_margin  # Inside
    section2.top_margin = section0.top_margin
    section2.bottom_margin = section0.bottom_margin
    section2.gutter = section0.gutter
    # Ensure mirrored margins are also copied
    sectPr2 = section2._sectPr
    pgMar2 = sectPr2.find(qn("w:pgMar"))
    if pgMar2 is not None:
        pgMar2.set(qn("w:mirrorMargins"), "true")

    # --- Section 2: Rest of Front Matter ---
    logging.info("Adding Rest of Front Matter (Section 2, starts page iii)...")
    fm_order = ["dedication", "foreword", "preface", "acknowledgements"]
    has_fm_content = False
    for key in fm_order:
        content = front_matter.get(key)
        # Check if content exists and is not the placeholder failure message
        if content and not content.startswith(
            f"[{key.title()} content generation failed.]"
        ):
            if has_fm_content:  # Add page break before subsequent FM sections
                doc.add_page_break()

            title = key.replace("_", " ").title()
            # Add space before title is handled by Heading 1 style now
            doc.add_paragraph(title, style="Heading 1")
            # Pass the main doc object as the container for markdown conversion
            context_label = f"FrontMatter_{key}"
            markdown_to_docx(
                content,
                doc,
                doc,
                config,
                usable_width_inches,
                equation_image_dir,
                context_label=context_label,
            )
            has_fm_content = True

    # Set page numbering for this section (starts after title/copyright)
    # Roman numerals starting from iii (page 3 conceptually)
    set_page_numbering(section2, format_code="lowerRoman", start_number=3)
    logging.info("Set Front Matter page numbering (lowerRoman, starting iii).")

    # --- Section Break for Body Matter (Starts Section 3) ---
    doc.add_section(WD_SECTION.ODD_PAGE)
    section3 = doc.sections[3]
    # Copy page setup from section 0 to section 3
    section3.page_height = section0.page_height
    section3.page_width = section0.page_width
    section3.left_margin = section0.left_margin  # Outside
    section3.right_margin = section0.right_margin  # Inside
    section3.top_margin = section0.top_margin
    section3.bottom_margin = section0.bottom_margin
    section3.gutter = section0.gutter
    # Ensure mirrored margins are also copied
    sectPr3 = section3._sectPr
    pgMar3 = sectPr3.find(qn("w:pgMar"))
    if pgMar3 is not None:
        pgMar3.set(qn("w:mirrorMargins"), "true")

    # --- Section 3: Body Matter (Starts page 1) ---
    logging.info("Adding Body Matter (Chapters)...")
    chapter_keys = list(body_matter.keys())
    for i, chapter_title in enumerate(chapter_keys):
        sections_data = body_matter[chapter_title]
        logging.info(f"Adding Chapter {i+1}: {chapter_title}")
        # Add chapter title (spacing handled by style)
        doc.add_paragraph(chapter_title, style="Heading 1")

        for j, section_info in enumerate(sections_data):
            section_content = section_info.get("content", "[Missing Content]")
            section_title = section_info.get("title", f"Section {j+1}")
            logging.debug(f"Adding content for Section {j+1}: '{section_title}'")
            # Add section title (spacing handled by style)
            if section_title:
                doc.add_paragraph(section_title, style="Heading 2")

            # --- Construct context label for body section ---
            safe_chap_title = sanitize_filename(chapter_title, 30)
            safe_sec_title = sanitize_filename(section_title, 30)
            context_label = f"Chap{i+1}_{safe_chap_title}_Sec{j+1}_{safe_sec_title}"
            # --- End context label construction ---

            # Pass the main doc object as the container
            markdown_to_docx(
                section_content,
                doc,
                doc,
                config,
                usable_width_inches,
                equation_image_dir,
                context_label=context_label,
            )

        # Add page break after chapter, except for the last one
        if i < len(chapter_keys) - 1:
            doc.add_page_break()

    # Set page numbering for the body section (starts at 1)
    set_page_numbering(section3, format_code="decimal", start_number=1)
    logging.info("Set Body Matter page numbering (decimal, starting at 1).")

    # --- Section Break for Back Matter (Starts Section 4) ---
    doc.add_section(WD_SECTION.ODD_PAGE)
    section4 = doc.sections[4]
    # Copy page setup from section 0 to section 4
    section4.page_height = section0.page_height
    section4.page_width = section0.page_width
    section4.left_margin = section0.left_margin  # Outside
    section4.right_margin = section0.right_margin  # Inside
    section4.top_margin = section0.top_margin
    section4.bottom_margin = section0.bottom_margin
    section4.gutter = section0.gutter
    # Ensure mirrored margins are also copied
    sectPr4 = section4._sectPr
    pgMar4 = sectPr4.find(qn("w:pgMar"))
    if pgMar4 is not None:
        pgMar4.set(qn("w:mirrorMargins"), "true")

    # --- Section 4: Back Matter (Continues numbering) ---
    logging.info("Adding Back Matter (Section 4, continuing numbering)...")
    bm_order = ["appendix", "glossary", "bibliography", "about_the_author"]

    # Conditionally exclude "About the Author" from the main book based on config
    gen_params = config.get("generation_params", {})
    include_about_author_in_main = gen_params.get(
        "include_about_author_in_main_book", True
    )  # Default to True

    if not include_about_author_in_main:
        if "about_the_author" in bm_order:
            bm_order.remove("about_the_author")
            logging.info(
                "'About the Author' section will be excluded from the main book as per config."
            )

    has_bm_content = False
    bm_added_count = 0
    valid_bm_keys = [
        key
        for key in bm_order
        if back_matter.get(key)
        and (
            (
                isinstance(back_matter[key], list) and key == "appendix"
            )  # Appendix is valid if it's a list
            or (
                isinstance(
                    back_matter[key], str
                )  # For strings (including appendix placeholders or other items)
                and not back_matter[key].startswith(
                    f"[{key.replace('_', ' ').title()} content generation failed.]"
                )
            )
        )
    ]

    for i, key in enumerate(valid_bm_keys):
        content = back_matter[key]
        if has_bm_content:  # Add page break before subsequent BM sections
            doc.add_page_break()
        title = key.replace("_", " ").title()
        # Add title (spacing handled by style)
        doc.add_paragraph(title, style="Heading 1")

        if key == "appendix" and isinstance(content, list):
            # Content is a list of appendix subsections
            num_subsections = len(content)
            for idx, subsection_data in enumerate(content):
                sub_title = subsection_data.get("title", "Untitled Subsection")
                sub_content_md = subsection_data.get("content", "[Content missing]")

                doc.add_paragraph(sub_title, style="Heading 2")  # Add subsection title
                context_label_sub = (
                    f"BackMatter_Appendix_{sanitize_filename(sub_title, 30)}"
                )
                markdown_to_docx(
                    sub_content_md,
                    doc,  # container_obj
                    doc,  # doc (main document)
                    config,
                    usable_width_inches,
                    equation_image_dir,
                    context_label=context_label_sub,
                )
                if idx < num_subsections - 1:  # If not the last subsection
                    doc.add_page_break()

        else:
            # Existing behavior for other back matter items (content is a string)
            # or if appendix is a fallback string placeholder
            context_label = f"BackMatter_{key}"
            markdown_to_docx(
                (
                    content
                    if isinstance(content, str)
                    else "[Invalid content format for this back matter item]"
                ),
                doc,  # container_obj
                doc,  # doc (main document)
                config,
                usable_width_inches,
                equation_image_dir,
                context_label=context_label,
            )

        has_bm_content = True
        bm_added_count += 1

    # Set page numbering for this section (continues from body)
    # Pass None for start_number to continue sequence
    set_page_numbering(section4, format_code="decimal", start_number=None)
    logging.info("Set Back Matter page numbering (continuing decimal).")

    # --- Save Document ---
    try:
        doc.save(output_filename)
        logging.info(f"Main DOCX file assembly complete. Saved as '{output_filename}'")
        return output_filename
    except PermissionError:
        logging.error(
            f"PermissionError: Could not save '{output_filename}'. Check if the file is open or permissions are correct."
        )
        return None  # Indicate failure
    except Exception as e:
        logging.error(f"Error saving main DOCX file '{output_filename}': {e}")
        return None  # Indicate failure


# --- Marketing Docx Assembly (using python-docx) ---
def assemble_marketing_docx(
    config,
    back_matter_content,
    blurb_content,
    summary_context,
    main_book_filename_stem,
    equation_image_dir,
    output_dir,
):
    """
    Assembles the separate marketing DOCX file using python-docx.
    """
    if not main_book_filename_stem:
        logging.error("Cannot create marketing docx without main book filename stem.")
        return
    # Construct the full output path using the provided directory
    output_filename = (
        output_dir / f"{main_book_filename_stem}_Marketing.docx"
    )  # Use pathlib's / operator
    logging.info(f"Assembling marketing DOCX file: '{output_filename}'")
    doc = Document()
    style_config = config.get("style_params", {})
    font_name = style_config.get("font_name", "Times New Roman")
    font_size = style_config.get("font_size", 12)
    gen_params = config.get("generation_params", {})

    try:  # Basic style setup
        style = doc.styles["Normal"]
        style.font.name = font_name
        style.font.size = Pt(font_size)

        # Define Heading 1 if not present
        if "Heading 1" not in doc.styles:
            h1_style = doc.styles.add_style("Heading 1", WD_STYLE_TYPE.PARAGRAPH)
            h1_style.base_style = doc.styles["Normal"]
            h1_style.font.name = font_name
            h1_style.font.size = Pt(16)
            h1_style.font.bold = True
            h1_style.paragraph_format.space_before = Pt(12)
            h1_style.paragraph_format.space_after = Pt(6)
        else:  # Ensure font consistency
            doc.styles["Heading 1"].font.name = font_name

        # Define Heading 2 if not present
        if "Heading 2" not in doc.styles:
            h2_style = doc.styles.add_style("Heading 2", WD_STYLE_TYPE.PARAGRAPH)
            h2_style.base_style = doc.styles["Normal"]
            h2_style.font.name = font_name
            h2_style.font.size = Pt(13)
            h2_style.font.bold = True
            h2_style.paragraph_format.space_before = Pt(10)
            h2_style.paragraph_format.space_after = Pt(4)
        else:  # Ensure font consistency
            doc.styles["Heading 2"].font.name = font_name

        # Define List Bullet if not present
        if "List Bullet" not in doc.styles:
            lb_style = doc.styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)
            lb_style.base_style = doc.styles["Normal"]
            # Add basic bullet formatting if needed

    except Exception as e:
        logging.warning(f"Could not apply basic styles to marketing doc: {e}")

    # --- Add Book Details Section ---
    doc.add_paragraph("Book Details", style="Heading 1")

    random_topic_seed = gen_params.get("random_topic_seed")

    book_title_marketing = gen_params.get("book_title", "[Not Specified]")
    doc.add_paragraph("Book Title:", style="Heading 2")
    doc.add_paragraph(book_title_marketing)

    book_subtitle_marketing = gen_params.get("book_subtitle", "[Not Specified]")
    doc.add_paragraph("Book Subtitle:", style="Heading 2")
    doc.add_paragraph(book_subtitle_marketing)
    if random_topic_seed:
        doc.add_paragraph("Random Topic Seed:", style="Heading 2")
        doc.add_paragraph(random_topic_seed)

    api_settings_local = config.get(
        "api_settings", {}
    )  # Use a local var to avoid conflict
    api_provider_local = api_settings_local.get("provider", "gemini")

    model_name_used = ""
    if api_provider_local == "gemini":
        model_name_used = api_settings_local.get("gemini", {}).get("model", "")
    elif api_provider_local == "ollama":
        model_name_used = api_settings_local.get("ollama", {}).get("model", "")
    else:
        model_name_used = "[Unknown API Provider or Model]"
        logging.warning(
            f"Unknown API provider '{api_provider_local}' when trying to get model name for marketing doc."
        )

    doc.add_paragraph("LLM Model Used:", style="Heading 2")
    doc.add_paragraph(model_name_used)

    doc.add_paragraph("Main Topic:", style="Heading 2")
    doc.add_paragraph(gen_params.get("main_topic", "[Not Specified]"))

    writing_tone = gen_params.get("writing_tone", DEFAULT_WRITING_TONE)
    doc.add_paragraph("Writing Tone:", style="Heading 2")
    doc.add_paragraph(writing_tone)

    doc.add_paragraph("Setting:", style="Heading 2")
    doc.add_paragraph(gen_params.get("setting", "[Not Specified]"))

    # --- Add Character List (if generated) ---
    character_list = gen_params.get("character_list")
    if character_list:
        doc.add_paragraph("characters", style="Heading 2")
        list_style = (
            doc.styles["List Bullet"]
            if "List Bullet" in doc.styles
            else doc.styles["Normal"]
        )
        for char in character_list:
            if isinstance(char, dict) and "name" in char and "description" in char:
                # Add name in bold, then description
                p = doc.add_paragraph(style=list_style)
                p.add_run(f"{char['name']}: ").bold = True
                p.add_run(char["description"])
            else:  # Fallback for unexpected format
                doc.add_paragraph(str(char), style=list_style)
    # --- End Character List ---

    doc.add_paragraph("Author Details:", style="Heading 2")
    doc.add_paragraph(f"Name: {gen_params.get('author_name', '[Not Specified]')}")
    doc.add_paragraph(f"Gender: {gen_params.get('author_gender', '[Not Specified]')}")

    doc.add_page_break()

    # --- Add Book Blurb ---
    doc.add_paragraph("Book Blurb", style="Heading 1")
    markdown_to_docx(
        blurb_content or "[Blurb generation failed]",
        doc,
        doc,
        config,
        None,
        equation_image_dir,
        context_label="Marketing_Blurb",
    )
    doc.add_page_break()

    # --- Add Book Summary Section ---
    doc.add_paragraph("Book Summary (from Chapter Summaries)", style="Heading 1")
    cleaned_summary = (summary_context or "").strip()
    if cleaned_summary and cleaned_summary != "[No chapter summaries available]":
        # Remove the initial label if present
        cleaned_summary = re.sub(
            r"^\s*Chapter summaries:\s*", "", cleaned_summary, flags=re.IGNORECASE
        ).strip()
        markdown_to_docx(
            cleaned_summary,
            doc,
            doc,
            config,
            None,
            equation_image_dir,
            context_label="Marketing_Summary",
        )
    else:
        doc.add_paragraph("[No summary points available]")
    doc.add_page_break()

    # --- Add About the Author ---
    doc.add_paragraph("About the Author", style="Heading 1")
    about_author_content = back_matter_content.get(
        "about_the_author", "[About the Author generation failed]"
    )
    markdown_to_docx(
        about_author_content,
        doc,
        doc,
        config,
        None,
        equation_image_dir,
        context_label="Marketing_AboutAuthor",
    )

    # --- Save the document ---
    try:
        doc.save(output_filename)
        logging.info(
            f"Marketing DOCX file assembly complete. Saved as '{output_filename}'"
        )
    except PermissionError:
        logging.error(
            f"PermissionError: Could not save marketing file '{output_filename}'. Check if file is open or permissions are correct."
        )
    except Exception as e:
        logging.error(f"Error saving marketing DOCX file '{output_filename}': {e}")
